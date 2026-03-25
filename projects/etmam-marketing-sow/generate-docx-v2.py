#!/usr/bin/env python3
"""Generate Pyramedia X proposal DOCX v2 — professional Arabic RTL."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import copy

doc = Document()

ORANGE = RGBColor(249, 115, 22)
BLACK = RGBColor(0x11, 0x11, 0x11)
DARK = RGBColor(0x1a, 0x1a, 0x1a)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT = 'Sakkal Majalla'  # Arabic font available in Windows
FONT_ALT = 'Arial'

# ============ HELPERS ============

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            el.set(qn('w:space'), '0')
            tcBorders.append(el)
    tcPr.append(tcBorders)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def add_run(paragraph, text, bold=False, color=None, size=None, font=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font or FONT
    run.font.size = Pt(size or 12)
    if color:
        run.font.color.rgb = color
    # Set complex script font for Arabic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), font or FONT)
    # Set complex script size
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str((size or 12) * 2))
    rPr.append(szCs)
    # Set complex script bold
    if bold:
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)
    set_rtl_run(run)
    return run

def make_para(align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    set_rtl(p)
    return p

def add_heading_orange(text, level=1):
    p = make_para(space_before=12 if level == 1 else 8, space_after=8)
    if level == 1:
        add_run(p, text, bold=True, color=BLACK, size=22)
        # Add orange bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:color'), 'F97316')
        bottom.set(qn('w:space'), '4')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        add_run(p, '◼ ', bold=True, color=ORANGE, size=14)
        add_run(p, text, bold=True, color=ORANGE, size=14)
    elif level == 3:
        add_run(p, text, bold=True, color=ORANGE, size=12)
    return p

def add_text(text, bold=False, color=None, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = make_para(align=align)
    add_run(p, text, bold=bold, color=color or DARK, size=size)
    return p

def add_numbered_item(num, text, start=None):
    p = make_para(space_after=4)
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, f'  {num}  ', bold=True, color=WHITE, size=10)
    # Fake the orange circle with just the number
    add_run(p, f'  {text}', color=DARK, size=12)
    return p

def add_styled_numbered(num, text):
    """Numbered item with orange number."""
    p = make_para(space_after=4)
    add_run(p, f'{num}. ', bold=True, color=ORANGE, size=12)
    add_run(p, text, color=DARK, size=12)
    return p

def add_bullet_item(text, prefix=None):
    p = make_para(space_after=3)
    p.paragraph_format.left_indent = Cm(0.8)
    if prefix:
        add_run(p, f'{prefix} ', bold=True, color=DARK, size=11)
    else:
        add_run(p, '• ', bold=True, color=ORANGE, size=12)
    add_run(p, text, color=DARK, size=11)
    return p

def add_note(text):
    p = make_para(space_before=6, space_after=10)
    # Add right border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    right_bdr = OxmlElement('w:right')
    right_bdr.set(qn('w:val'), 'single')
    right_bdr.set(qn('w:sz'), '12')
    right_bdr.set(qn('w:color'), 'F97316')
    right_bdr.set(qn('w:space'), '8')
    pBdr.append(right_bdr)
    pPr.append(pBdr)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, '📝 ملاحظة: ', bold=True, color=ORANGE, size=10)
    add_run(p, text, color=GRAY, size=10, italic=True)
    return p

def add_page_break():
    doc.add_page_break()

def make_table(rows, cols, data, header_row=True, col_widths=None):
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style all cells
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Set padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom', 'start', 'end']:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), '80')
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
            
            if col_widths and c_idx < len(col_widths):
                set_cell_width(cell, col_widths[c_idx])
            
            if r_idx == 0 and header_row:
                set_cell_shading(cell, 'F97316')
                set_cell_borders(cell, 
                    top={'color': 'F97316', 'sz': 6},
                    bottom={'color': 'F97316', 'sz': 6},
                    left={'color': 'F97316', 'sz': 4},
                    right={'color': 'F97316', 'sz': 4})
            else:
                border_color = 'EEEEEE'
                set_cell_borders(cell,
                    bottom={'color': border_color, 'sz': 4},
                    left={'color': border_color, 'sz': 2},
                    right={'color': border_color, 'sz': 2})
                if r_idx % 2 == 0 and r_idx > 0:
                    set_cell_shading(cell, 'FAFAFA')
    
    # Fill data
    for r_idx, row_data in enumerate(data):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            for para in cell.paragraphs:
                para.clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(p)
            
            if r_idx == 0 and header_row:
                add_run(p, str(cell_text), bold=True, color=WHITE, size=11)
            else:
                is_bold = r_idx == 2 and header_row  # Recommended row
                add_run(p, str(cell_text), bold=is_bold, color=DARK, size=10)
    
    return table

# ============ PAGE SETUP ============
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    # RTL document
    sectPr = section._sectPr
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    sectPr.append(bidi)

# Set default font
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(12)
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:cs'), FONT)
rPr.append(rFonts)

# ==================== COVER ====================
for _ in range(5):
    doc.add_paragraph()

p = make_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_run(p, 'PYRAMEDIA', bold=True, color=BLACK, size=34, font='Arial')
add_run(p, 'X', bold=True, color=ORANGE, size=34, font='Arial')

p = make_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
add_run(p, 'FOR AI SOLUTIONS', color=ORANGE, size=12, font='Arial')

p = make_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
# Orange line
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '18')
bottom.set(qn('w:color'), 'F97316')
bottom.set(qn('w:space'), '1')
pBdr.append(bottom)
pPr.append(pBdr)

doc.add_paragraph()

p = make_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_run(p, 'عرض خدمات التسويق الرقمي الشامل', bold=True, color=BLACK, size=26)

p = make_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
add_run(p, 'مركز إتمام للخدمات القضائية', bold=True, color=ORANGE, size=20)

for _ in range(3):
    doc.add_paragraph()

add_text('مارس 2026', color=GRAY, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text('مقدم من: Pyramedia X — For AI Solutions', color=GRAY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_text('سري وخاص — هذا العرض مقدم حصرياً للمعني بالأمر', color=LIGHT_GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

# ==================== PAGE 2: INTRO ====================
add_page_break()
add_heading_orange('المقدمة')

for text in [
    'تتشرّف شركة Pyramedia X بتقديم هذا العرض الشامل لخدمات التسويق الرقمي وإدارة المحتوى والعلامة التجارية لصالح مركز إتمام للخدمات القضائية. نضع في هذا العرض خبراتنا المتراكمة في التسويق الرقمي وإنتاج المحتوى الاحترافي وإدارة الحملات الإعلانية في خدمة المركز لتعزيز حضوره الرقمي وبناء سمعة قوية تليق بمكانته.',
    'نؤمن بأن المؤسسات القضائية والقانونية تستحق تسويقاً رقمياً يعكس هيبتها ومصداقيتها، ولهذا صُمّمت خدماتنا لتراعي الطابع المهني والاحترافي مع الوصول الفعّال للجمهور المستهدف عبر مختلف القنوات الرقمية.',
    'يتضمن هذا العرض تفاصيل نطاق العمل، الكميات، التسعير، والشروط التعاقدية الكاملة.',
]:
    p = make_para(space_after=10)
    p.paragraph_format.line_spacing = 1.8
    add_run(p, text, color=RGBColor(0x33, 0x33, 0x33), size=12)

# ==================== PAGE 3: SETUP ====================
add_page_break()
add_heading_orange('نطاق العمل — مهام التأسيس (Setup)')
add_text('مهام تُنفّذ مرة واحدة في بداية التعاقد لتهيئة البنية التسويقية الرقمية الكاملة للمركز.', color=GRAY, size=11)

setup_items = [
    'إنشاء وتأسيس الحسابات الرسمية على: Instagram, Facebook, TikTok, LinkedIn',
    'تطبيق الهوية البصرية على جميع المنصات (قوالب، أيقونات، أغلفة، الوصف التعريفي)',
    'إعداد الأعمدة التحريرية (Content Pillars) والخطوط التحريرية',
    'إعداد حسابات الإعلانات (Meta Business Suite, Google Ads)',
    'إعداد Google Business Profile',
    'بناء الخطة التسويقية الأولى (أول 3 أشهر)',
    'تحليل المنافسين الأولي',
    'تحديد الجمهور المستهدف وبناء شخصيات العملاء (Personas)',
    'بناء خطة إدارة الأزمات (Crisis Communication Protocol)',
]
for i, item in enumerate(setup_items, 1):
    add_styled_numbered(i, item)

doc.add_paragraph()
p = make_para(space_before=8)
pPr = p._p.get_or_add_pPr()
# Add shading background
shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF7ED" w:val="clear"/>')
pPr.append(shd)
pBdr = OxmlElement('w:pBdr')
for side in ['top', 'bottom', 'left', 'right']:
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '6')
    b.set(qn('w:color'), 'F97316')
    b.set(qn('w:space'), '8')
    pBdr.append(b)
pPr.append(pBdr)
add_run(p, '⏱ مدة التأسيس: ', bold=True, color=ORANGE, size=12)
add_run(p, 'من 17 إلى 25 يوم عمل من تاريخ بدء التعاقد', color=DARK, size=12)

p = make_para()
pPr = p._p.get_or_add_pPr()
shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF7ED" w:val="clear"/>')
pPr.append(shd)
pBdr = OxmlElement('w:pBdr')
for side in ['top', 'bottom', 'left', 'right']:
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '6')
    b.set(qn('w:color'), 'F97316')
    b.set(qn('w:space'), '8')
    pBdr.append(b)
pPr.append(pBdr)
add_run(p, '💰 تكلفة التأسيس: ', bold=True, color=ORANGE, size=12)
add_run(p, 'مشمولة بالكامل ضمن قيمة العقد', color=DARK, size=12)

# ==================== PAGE 4: MONTHLY SOCIAL + CONTENT ====================
add_page_break()
add_heading_orange('نطاق العمل — المهام الشهرية المستمرة')

add_heading_orange('أولاً: إدارة منصات التواصل الاجتماعي', level=2)
social = [
    'إعداد تقويم محتوى شهري معتمد',
    'كتابة وتصميم ونشر المحتوى على 4 منصات (Instagram, Facebook, TikTok, LinkedIn)',
    'إدارة التعليقات والرسائل على مدار الساعة (24/7)',
    'متابعة التريندات والمناسبات القانونية ذات الصلة',
]
for i, item in enumerate(social, 1):
    add_styled_numbered(i, item)

add_heading_orange('ثانياً: إنتاج المحتوى', level=2)
content = [
    'كتابة سكريبتات الفيديو (ريلز / تيكتوك)',
    'تصوير ومونتاج المحتوى المرئي',
    'جلسات تصوير احترافية — جلستين شهرياً (يوم كامل لكل جلسة) تشمل: تصوير فريق العمل، المقر، ومحتوى السوشيال ميديا',
    'تصميم الجرافيك (بوستات + ستوريز + كاروسيل)',
    'كتابة المحتوى باللغتين العربية والإنجليزية',
]
for i, item in enumerate(content, 5):
    add_styled_numbered(i, item)

add_note('يتم توفير المحتوى بلغات إضافية (الأوردو / الهندية) كميزة تعزيزية للمنشورات والفيديوهات عند ملاءمة الفكرة والسياق، دون أن يشكّل ذلك التزاماً تعاقدياً.')

# ==================== PAGE 5: QUANTITIES + ADS ====================
add_page_break()
add_heading_orange('الكميات الشهرية', level=2)

qty_data = [
    ['البند', 'الكمية'],
    ['منشورات (بوست + ريلز + كاروسيل)', '28 منشور/شهر'],
    ['ستوريز', 'يومياً'],
    ['مقالات قانونية للموقع الإلكتروني', '6 مقالات/شهر'],
    ['فيديوهات قانونية (ريلز) ضمن الـ 28 منشور', 'حتى 5 فيديوهات/شهر'],
    ['جلسات تصوير احترافية', 'جلستين/شهر (يوم كامل)'],
]
make_table(6, 2, qty_data, col_widths=[11, 6])

add_note('توزيع الـ 28 منشور بين بوستات وريلز وكاروسيل يكون مرناً حسب الخطة الشهرية واحتياجات المحتوى. الفيديوهات القانونية الخمسة هي ضمن الـ 28 منشور وليست إضافية.')

add_heading_orange('ثالثاً: الإعلانات المدفوعة', level=2)
ads = [
    'إدارة حملات Meta (Instagram + Facebook)',
    'إدارة حملات Google Ads',
    'استهداف وتحسين مستمر واختبارات A/B Testing',
    'إعداد تقارير أداء الحملات',
]
for i, item in enumerate(ads, 10):
    add_styled_numbered(i, item)

add_note('رسوم إدارة الحملات الإعلانية مشمولة بالكامل ضمن قيمة العقد. الميزانية الإعلانية تكون على حساب العميل مباشرةً.')

# ==================== PAGE 6: SEO + REPUTATION + STRATEGY ====================
add_page_break()

add_heading_orange('رابعاً: تحسين محركات البحث (SEO) وتسويق المحتوى', level=2)
for i, item in enumerate(['تحسين SEO للموقع الإلكتروني (بعد تسليمه من العميل)', 'كتابة مقالات قانونية تعليمية (6 مقالات شهرياً)', 'تحسين Local SEO'], 14):
    add_styled_numbered(i, item)

add_heading_orange('خامساً: إدارة السمعة والعلامة التجارية', level=2)
for i, item in enumerate(['متابعة وإدارة تقييمات Google Reviews', 'تشجيع العملاء على التقييم', 'مراقبة العلامة التجارية أونلاين (Mentions + تنبيهات)', 'جمع شهادات العملاء (تصوير / كتابة قصص نجاح)'], 17):
    add_styled_numbered(i, item)

add_heading_orange('سادساً: التخطيط الاستراتيجي', level=2)
for i, item in enumerate(['تحديث الخطة التسويقية كل ربع سنة', 'تحليل منافسين دوري', 'اقتراح حملات موسمية ومبادرات', 'تحديث خطة الأزمات ربع سنوي'], 21):
    add_styled_numbered(i, item)

add_heading_orange('سابعاً: التقارير والمتابعة', level=2)
for i, item in enumerate(['تقرير أداء شهري شامل', 'اجتماع شهري مع الإدارة', 'متابعة مؤشرات الأداء (KPIs)'], 25):
    add_styled_numbered(i, item)

add_heading_orange('ثامناً: التنسيق والإشراف', level=2)
add_styled_numbered(28, 'التنسيق بين مركز إتمام وموردي الخدمات (مطبوعات، فعاليات، مؤتمرات) والإشراف على جودة التنفيذ بما يضمن تقديم المركز بالصورة اللائقة بدائرة القضاء')

# ==================== PAGE 7: EXTRA VIDEOS + SERVICES ====================
add_page_break()
add_heading_orange('الفيديوهات القانونية الإضافية')
add_text('يشمل العقد إنتاج حتى 5 فيديوهات قانونية (ريلز) شهرياً ضمن الـ 28 منشور.', size=12)
add_text('أي فيديو إضافي يُسعّر على النحو التالي:', size=12)

doc.add_paragraph()
p = make_para(space_before=8, space_after=8)
pPr = p._p.get_or_add_pPr()
shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF7ED" w:val="clear"/>')
pPr.append(shd)
add_run(p, '💰 سعر الفيديو القانوني الإضافي: 650 درهم', bold=True, color=ORANGE, size=14)

add_bullet_item('كتابة السكريبت', prefix='✅')
add_bullet_item('تدريب المتحدث', prefix='✅')
add_bullet_item('التصوير', prefix='✅')
add_bullet_item('المونتاج', prefix='✅')
add_bullet_item('غير شامل إيجار الاستوديو في حالة التصوير خارج مقر المركز', prefix='❌')

doc.add_paragraph()
add_heading_orange('خدمات إضافية (حسب الطلب — بتسعير منفصل)')
add_text('يمكن طلب أي من الخدمات التالية بتسعير يُحدد بناءً على متطلبات كل مشروع على حدة:', color=GRAY, size=11)

for cat, items in [
    ('الإنتاج', ['تغطية فعاليات ومؤتمرات', 'إنتاج فيديوهات مؤسسية (Brand Film)', 'موشن جرافيك']),
    ('التسويق', ['حملات موسمية خاصة', 'علاقات عامة وإعلامية (PR)', 'التسويق عبر المؤثرين']),
    ('التسويق التقليدي', ['مطبوعات (بروشورات، فلايرز، بانرات)', 'تصميم أكشاك ومعارض', 'لوحات إعلانية خارجية']),
    ('الخدمات الرقمية', ['تصميم صفحات هبوط (Landing Pages)', 'إنتاج بودكاست قانوني']),
]:
    add_heading_orange(cat, level=3)
    for item in items:
        add_bullet_item(item)

# ==================== PAGE 8: PRICING ====================
add_page_break()
add_heading_orange('التسعير')

pricing_data = [
    ['خيار الدفع', 'السعر الشهري', 'إجمالي الفترة', 'التوفير السنوي'],
    ['ربع سنوي\n(كل 3 أشهر)', '18,000\nدرهم', '54,000\nدرهم', '—'],
    ['⭐ نصف سنوي\n(كل 6 أشهر)', '16,000\nدرهم', '96,000\nدرهم', 'توفير\n12,000 درهم'],
    ['سنوي\n(دفعة واحدة)', '15,000\nدرهم', '180,000\nدرهم', 'توفير\n36,000 درهم'],
]
table = make_table(4, 4, pricing_data, col_widths=[5, 4, 4, 4])

# Highlight recommended row
for cell in table.rows[2].cells:
    set_cell_shading(cell, 'FFF7ED')

add_text('* جميع الأسعار بالدرهم الإماراتي وغير شاملة ضريبة القيمة المضافة (إن وُجدت)', color=LIGHT_GRAY, size=9)

doc.add_paragraph()
add_heading_orange('يشمل السعر', level=3)
for item in ['جميع مهام التأسيس (Setup) — 9 بنود', 'جميع المهام الشهرية المستمرة — 28 بنداً', 'إدارة الحملات الإعلانية على Meta وGoogle', 'إنتاج حتى 5 فيديوهات قانونية شهرياً']:
    add_bullet_item(item, prefix='✅')

doc.add_paragraph()
add_heading_orange('لا يشمل السعر', level=3)
for item in ['الميزانية الإعلانية (على حساب العميل مباشرةً)', 'الخدمات الإضافية (بتسعير منفصل حسب الطلب)', 'إيجار الاستوديو للفيديوهات القانونية الإضافية']:
    add_bullet_item(item, prefix='❌')

# ==================== PAGE 9: TERMS ====================
add_page_break()
add_heading_orange('الشروط والأحكام')

sections_terms = [
    ('مدة العقد', [
        'مدة العقد سنة واحدة تبدأ من تاريخ التوقيع',
        'يُجدد العقد تلقائياً ما لم يُخطر أحد الطرفين الآخر خطياً قبل 30 يوماً من انتهاء مدته',
    ]),
    ('آلية العمل والاعتماد', [
        'يُقدَّم تقويم المحتوى الشهري للاعتماد قبل بداية كل شهر',
        'يتم اعتماد المحتوى من قبل الشخص أو الأشخاص المخولين من طرف العميل (بحد أقصى شخصين)',
        'المراجعات والتعديلات على المحتوى مفتوحة بشرط تقديمها قبل 48 ساعة من موعد النشر المُحدد',
    ]),
    ('الدفع', [
        'يتم الدفع مقدماً حسب خيار الدفع المُختار (ربع سنوي / نصف سنوي / سنوي)',
        'في حالة التأخر عن السداد، يحق لبيراميديا تعليق جميع الخدمات حتى استلام المستحقات المالية كاملة',
    ]),
    ('الإلغاء', [
        'يحق لأي من الطرفين إنهاء العقد بإشعار خطي مسبق لا يقل عن 30 يوماً',
        'في حالة الإلغاء، يلتزم العميل بسداد جميع المستحقات عن الفترة المنقضية حتى تاريخ الإنهاء الفعلي',
    ]),
    ('الملكية الفكرية', [
        'جميع المحتوى والمواد الإبداعية المُنتجة خصيصاً للعميل تؤول ملكيتها بالكامل إلى العميل بعد سداد جميع المستحقات المالية',
        'تحتفظ بيراميديا بحقها في استخدام الأعمال المُنجزة ضمن معرض أعمالها (Portfolio) ما لم يُتفق على خلاف ذلك',
    ]),
    ('ملاحظات عامة', [
        'الهوية البصرية: تم تسليمها مسبقاً بعقد منفصل',
        'الموقع الإلكتروني: مسؤولية العميل — تعمل بيراميديا على تحسين SEO فقط بعد تسليم الموقع',
        'منصة المحامين: مشروع منفصل بعقد مستقل',
    ]),
]

for title, items in sections_terms:
    add_heading_orange(title, level=2)
    for item in items:
        add_bullet_item(item)

# ==================== PAGE 10: SIGNATURE ====================
add_page_break()
add_heading_orange('التوقيع والاعتماد')
add_text('بتوقيع الطرفين أدناه، يُعتبر هذا العرض اتفاقاً ملزماً ونافذاً بالشروط والأحكام المذكورة أعلاه.', color=GRAY, size=11)

doc.add_paragraph()
add_heading_orange('الطرف الأول — بيراميديا', level=2)
for field in ['الاسم', 'المسمى الوظيفي', 'التوقيع', 'التاريخ']:
    p = make_para(space_after=14)
    add_run(p, f'{field}: ', bold=True, color=GRAY, size=11)
    add_run(p, '___________________________________', color=LIGHT_GRAY, size=11)

doc.add_paragraph()
add_heading_orange('الطرف الثاني — مركز إتمام للخدمات القضائية', level=2)
for field in ['الاسم', 'المسمى الوظيفي', 'رقم التواصل', 'التوقيع', 'التاريخ']:
    p = make_para(space_after=14)
    add_run(p, f'{field}: ', bold=True, color=GRAY, size=11)
    add_run(p, '___________________________________', color=LIGHT_GRAY, size=11)

doc.add_paragraph()
add_heading_orange('أشخاص اعتماد المحتوى (بحد أقصى شخصين)', level=2)
approval_data = [
    ['#', 'الاسم', 'رقم التواصل'],
    ['1', '', ''],
    ['2', '', ''],
]
make_table(3, 3, approval_data, col_widths=[2, 8, 7])

doc.add_paragraph()
add_heading_orange('خيار الدفع المُختار', level=2)
for opt in ['☐  ربع سنوي (18,000 درهم/شهر)', '☐  نصف سنوي (16,000 درهم/شهر)', '☐  سنوي (15,000 درهم/شهر)']:
    add_text(opt, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3):
    doc.add_paragraph()

add_text('سري وخاص — Pyramedia X © 2026', color=LIGHT_GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============ SAVE ============
out = '/home/node/openclaw/projects/etmam-marketing-sow/proposal.docx'
doc.save(out)
print(f'✅ DOCX v2 saved: {out}')
