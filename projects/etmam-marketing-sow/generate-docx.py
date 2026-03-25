#!/usr/bin/env python3
"""Generate Pyramedia X proposal DOCX for Itmam Marketing SOW."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# -- Page setup (A4, RTL margins)
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

ORANGE = RGBColor(249, 115, 22)
BLACK = RGBColor(0x11, 0x11, 0x11)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

style = doc.styles['Normal']
style.font.name = 'Tajawal'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.line_spacing = 1.6
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# RTL helper
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = BLACK
        run.bold = True
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = ORANGE
        run.bold = True
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = ORANGE
        run.bold = True
    p.paragraph_format.space_after = Pt(8)
    return p

def add_para(text, bold=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align or WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Clear any default runs
    for r in p.runs:
        r.text = ''
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(' ' + text)
    else:
        p.add_run(text)
    return p

def add_numbered(num, text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_num = p.add_run(f'{num}. ')
    run_num.bold = True
    run_num.font.color.rgb = ORANGE
    p.add_run(text)
    return p

def add_note(text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'📝 ملاحظة: {text}')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_page_break():
    doc.add_page_break()

# ==================== COVER ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
for _ in range(6):
    p.add_run('\n')

run = p.add_run('PYRAMEDIA')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = BLACK
run = p.add_run('X')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = ORANGE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p2)
run = p2.add_run('FOR AI SOLUTIONS')
run.font.size = Pt(12)
run.font.color.rgb = ORANGE

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run('ـــــــــــــــــــ')
run.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run('عرض خدمات التسويق الرقمي الشامل')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = BLACK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run('مركز إتمام للخدمات القضائية')
run.font.size = Pt(20)
run.font.color.rgb = ORANGE
run.bold = True

for _ in range(3):
    doc.add_paragraph()

add_para('مارس 2026', color=GRAY, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('مقدم من: Pyramedia X — For AI Solutions', color=GRAY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3):
    doc.add_paragraph()

add_para('سري وخاص — هذا العرض مقدم حصرياً للمعني بالأمر', color=RGBColor(0xBB, 0xBB, 0xBB), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

# ==================== PAGE 2: INTRO ====================
add_page_break()
add_heading_styled('المقدمة')
add_para('تتشرّف شركة Pyramedia X بتقديم هذا العرض الشامل لخدمات التسويق الرقمي وإدارة المحتوى والعلامة التجارية لصالح مركز إتمام للخدمات القضائية. نضع في هذا العرض خبراتنا المتراكمة في التسويق الرقمي وإنتاج المحتوى الاحترافي وإدارة الحملات الإعلانية في خدمة المركز لتعزيز حضوره الرقمي وبناء سمعة قوية تليق بمكانته.')
add_para('نؤمن بأن المؤسسات القضائية والقانونية تستحق تسويقاً رقمياً يعكس هيبتها ومصداقيتها، ولهذا صُمّمت خدماتنا لتراعي الطابع المهني والاحترافي مع الوصول الفعّال للجمهور المستهدف عبر مختلف القنوات الرقمية.')
add_para('يتضمن هذا العرض تفاصيل نطاق العمل، الكميات، التسعير، والشروط التعاقدية الكاملة.')

# ==================== PAGE 3: SETUP ====================
add_page_break()
add_heading_styled('نطاق العمل — مهام التأسيس (Setup)')
add_para('مهام تُنفّذ مرة واحدة في بداية التعاقد لتهيئة البنية التسويقية الرقمية الكاملة للمركز.', color=GRAY, size=11)

setup_items = [
    'إنشاء وتأسيس الحسابات الرسمية على: Instagram, Facebook, TikTok, LinkedIn',
    'تطبيق الهوية البصرية على جميع المنصات (قوالب، أيقونات، أغلفة، الوصف التعريفي)',
    'إعداد الأعمدة التحريرية (Content Pillars) والخطوط التحريرية',
    'إعداد حسابات الإعلانات (Meta Business Suite, Google Ads)',
    'إعداد Google Business Profile',
    'بناء الخطة التسويقية الأولى (أول 3 أشهر)',
    'تحليل المنافسين الأولي',
    'تحديد الجمهور المستهدف وبناء شخصيات العملاء (Personas)',
    'بناء خطة إدارة الأزمات (Crisis Communication Protocol)',
]
for i, item in enumerate(setup_items, 1):
    add_numbered(i, item)

doc.add_paragraph()
add_para('⏱ مدة التأسيس: من 17 إلى 25 يوم عمل من تاريخ بدء التعاقد', bold=True)
add_para('💰 تكلفة التأسيس: مشمولة بالكامل ضمن قيمة العقد', bold=True)

# ==================== PAGE 4: MONTHLY - SOCIAL + CONTENT ====================
add_page_break()
add_heading_styled('نطاق العمل — المهام الشهرية المستمرة')

add_heading_styled('أولاً: إدارة منصات التواصل الاجتماعي', level=2)
monthly_social = [
    'إعداد تقويم محتوى شهري معتمد',
    'كتابة وتصميم ونشر المحتوى على 4 منصات (Instagram, Facebook, TikTok, LinkedIn)',
    'إدارة التعليقات والرسائل على مدار الساعة (24/7)',
    'متابعة التريندات والمناسبات القانونية ذات الصلة',
]
for i, item in enumerate(monthly_social, 1):
    add_numbered(i, item)

add_heading_styled('ثانياً: إنتاج المحتوى', level=2)
content_items = [
    'كتابة سكريبتات الفيديو (ريلز / تيكتوك)',
    'تصوير ومونتاج المحتوى المرئي',
    'جلسات تصوير احترافية — جلستين شهرياً (يوم كامل لكل جلسة) تشمل: تصوير فريق العمل، المقر، ومحتوى السوشيال ميديا',
    'تصميم الجرافيك (بوستات + ستوريز + كاروسيل)',
    'كتابة المحتوى باللغتين العربية والإنجليزية',
]
for i, item in enumerate(content_items, 5):
    add_numbered(i, item)

add_note('يتم توفير المحتوى بلغات إضافية (الأوردو / الهندية) كميزة تعزيزية للمنشورات والفيديوهات عند ملاءمة الفكرة والسياق، دون أن يشكّل ذلك التزاماً تعاقدياً.')

# ==================== PAGE 5: QUANTITIES ====================
add_page_break()
add_heading_styled('الكميات الشهرية', level=2)

table = doc.add_table(rows=6, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
headers = ['البند', 'الكمية']
quantities = [
    ('منشورات (بوست + ريلز + كاروسيل)', '28 منشور/شهر'),
    ('ستوريز', 'يومياً'),
    ('مقالات قانونية للموقع الإلكتروني', '6 مقالات/شهر'),
    ('فيديوهات قانونية (ريلز) ضمن الـ 28 منشور', 'حتى 5 فيديوهات/شهر'),
    ('جلسات تصوير احترافية', 'جلستين/شهر (يوم كامل)'),
]

for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'F97316')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(p)
        for r in p.runs:
            r.font.color.rgb = WHITE
            r.font.bold = True
            r.font.size = Pt(11)

for row_idx, (item, qty) in enumerate(quantities, 1):
    table.rows[row_idx].cells[0].text = item
    table.rows[row_idx].cells[1].text = qty
    for cell in table.rows[row_idx].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(p)
            for r in p.runs:
                r.font.size = Pt(10)

add_note('توزيع الـ 28 منشور بين بوستات وريلز وكاروسيل يكون مرناً حسب الخطة الشهرية واحتياجات المحتوى. الفيديوهات القانونية الخمسة هي ضمن الـ 28 منشور وليست إضافية.')

add_heading_styled('ثالثاً: الإعلانات المدفوعة', level=2)
ads_items = [
    'إدارة حملات Meta (Instagram + Facebook)',
    'إدارة حملات Google Ads',
    'استهداف وتحسين مستمر واختبارات A/B Testing',
    'إعداد تقارير أداء الحملات',
]
for i, item in enumerate(ads_items, 10):
    add_numbered(i, item)
add_note('رسوم إدارة الحملات الإعلانية مشمولة بالكامل ضمن قيمة العقد. الميزانية الإعلانية تكون على حساب العميل مباشرةً.')

# ==================== PAGE 6: SEO + REPUTATION + STRATEGY ====================
add_page_break()
add_heading_styled('رابعاً: تحسين محركات البحث (SEO) وتسويق المحتوى', level=2)
seo_items = [
    'تحسين SEO للموقع الإلكتروني (بعد تسليمه من العميل)',
    'كتابة مقالات قانونية تعليمية (6 مقالات شهرياً)',
    'تحسين Local SEO',
]
for i, item in enumerate(seo_items, 14):
    add_numbered(i, item)

add_heading_styled('خامساً: إدارة السمعة والعلامة التجارية', level=2)
rep_items = [
    'متابعة وإدارة تقييمات Google Reviews',
    'تشجيع العملاء على التقييم',
    'مراقبة العلامة التجارية أونلاين (Mentions + تنبيهات)',
    'جمع شهادات العملاء (تصوير / كتابة قصص نجاح)',
]
for i, item in enumerate(rep_items, 17):
    add_numbered(i, item)

add_heading_styled('سادساً: التخطيط الاستراتيجي', level=2)
strat_items = [
    'تحديث الخطة التسويقية كل ربع سنة',
    'تحليل منافسين دوري',
    'اقتراح حملات موسمية ومبادرات',
    'تحديث خطة الأزمات ربع سنوي',
]
for i, item in enumerate(strat_items, 21):
    add_numbered(i, item)

add_heading_styled('سابعاً: التقارير والمتابعة', level=2)
report_items = [
    'تقرير أداء شهري شامل',
    'اجتماع شهري مع الإدارة',
    'متابعة مؤشرات الأداء (KPIs)',
]
for i, item in enumerate(report_items, 25):
    add_numbered(i, item)

add_heading_styled('ثامناً: التنسيق والإشراف', level=2)
add_numbered(28, 'التنسيق بين مركز إتمام وموردي الخدمات (مطبوعات، فعاليات، مؤتمرات) والإشراف على جودة التنفيذ بما يضمن تقديم المركز بالصورة اللائقة بدائرة القضاء')

# ==================== PAGE 7: EXTRA VIDEOS + SERVICES ====================
add_page_break()
add_heading_styled('الفيديوهات القانونية الإضافية')
add_para('يشمل العقد إنتاج حتى 5 فيديوهات قانونية (ريلز) شهرياً ضمن الـ 28 منشور. أي فيديو إضافي يُسعّر على النحو التالي:')
add_para('💰 سعر الفيديو القانوني الإضافي: 650 درهم', bold=True)
add_bullet('كتابة السكريبت', bold_prefix='✅')
add_bullet('تدريب المتحدث', bold_prefix='✅')
add_bullet('التصوير', bold_prefix='✅')
add_bullet('المونتاج', bold_prefix='✅')
add_bullet('غير شامل إيجار الاستوديو في حالة التصوير خارج مقر المركز', bold_prefix='❌')

doc.add_paragraph()
add_heading_styled('خدمات إضافية (حسب الطلب — بتسعير منفصل)')
add_para('يمكن طلب أي من الخدمات التالية بتسعير يُحدد بناءً على متطلبات كل مشروع على حدة:', color=GRAY)

add_heading_styled('الإنتاج', level=3)
for s in ['تغطية فعاليات ومؤتمرات', 'إنتاج فيديوهات مؤسسية (Brand Film)', 'موشن جرافيك']:
    add_bullet(s)

add_heading_styled('التسويق', level=3)
for s in ['حملات موسمية خاصة', 'علاقات عامة وإعلامية (PR)', 'التسويق عبر المؤثرين']:
    add_bullet(s)

add_heading_styled('التسويق التقليدي', level=3)
for s in ['مطبوعات (بروشورات، فلايرز، بانرات)', 'تصميم أكشاك ومعارض', 'لوحات إعلانية خارجية']:
    add_bullet(s)

add_heading_styled('الخدمات الرقمية', level=3)
for s in ['تصميم صفحات هبوط (Landing Pages)', 'إنتاج بودكاست قانوني']:
    add_bullet(s)

# ==================== PAGE 8: PRICING ====================
add_page_break()
add_heading_styled('التسعير')

table = doc.add_table(rows=4, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

pricing_headers = ['خيار الدفع', 'السعر الشهري', 'إجمالي الفترة', 'التوفير السنوي']
pricing_data = [
    ('ربع سنوي (كل 3 أشهر)', '18,000 درهم', '54,000 درهم', '—'),
    ('نصف سنوي (كل 6 أشهر) ⭐', '16,000 درهم', '96,000 درهم', 'توفير 12,000 درهم'),
    ('سنوي (دفعة واحدة)', '15,000 درهم', '180,000 درهم', 'توفير 36,000 درهم'),
]

for i, h in enumerate(pricing_headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'F97316')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(p)
        for r in p.runs:
            r.font.color.rgb = WHITE
            r.font.bold = True
            r.font.size = Pt(10)

for row_idx, row_data in enumerate(pricing_data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(p)
            for r in p.runs:
                r.font.size = Pt(10)
                if row_idx == 2:
                    r.font.bold = True

add_para('* جميع الأسعار بالدرهم الإماراتي وغير شاملة ضريبة القيمة المضافة (إن وُجدت)', color=GRAY, size=9)

doc.add_paragraph()
add_para('✅ يشمل السعر:', bold=True, color=ORANGE)
add_bullet('جميع مهام التأسيس (Setup) — 9 بنود')
add_bullet('جميع المهام الشهرية المستمرة — 28 بنداً')
add_bullet('إدارة الحملات الإعلانية على Meta وGoogle')
add_bullet('إنتاج حتى 5 فيديوهات قانونية شهرياً')

doc.add_paragraph()
add_para('❌ لا يشمل السعر:', bold=True, color=RGBColor(0xCC, 0x00, 0x00))
add_bullet('الميزانية الإعلانية (على حساب العميل مباشرةً)')
add_bullet('الخدمات الإضافية (بتسعير منفصل حسب الطلب)')
add_bullet('إيجار الاستوديو للفيديوهات القانونية الإضافية')

# ==================== PAGE 9: TERMS ====================
add_page_break()
add_heading_styled('الشروط والأحكام')

add_heading_styled('مدة العقد', level=2)
add_bullet('مدة العقد سنة واحدة تبدأ من تاريخ التوقيع')
add_bullet('يُجدد العقد تلقائياً ما لم يُخطر أحد الطرفين الآخر خطياً قبل 30 يوماً من انتهاء مدته')

add_heading_styled('آلية العمل والاعتماد', level=2)
add_bullet('يُقدَّم تقويم المحتوى الشهري للاعتماد قبل بداية كل شهر')
add_bullet('يتم اعتماد المحتوى من قبل الشخص أو الأشخاص المخولين من طرف العميل (بحد أقصى شخصين)')
add_bullet('المراجعات والتعديلات على المحتوى مفتوحة بشرط تقديمها قبل 48 ساعة من موعد النشر المُحدد')

add_heading_styled('الدفع', level=2)
add_bullet('يتم الدفع مقدماً حسب خيار الدفع المُختار (ربع سنوي / نصف سنوي / سنوي)')
add_bullet('في حالة التأخر عن السداد، يحق لبيراميديا تعليق جميع الخدمات حتى استلام المستحقات المالية كاملة')

add_heading_styled('الإلغاء', level=2)
add_bullet('يحق لأي من الطرفين إنهاء العقد بإشعار خطي مسبق لا يقل عن 30 يوماً')
add_bullet('في حالة الإلغاء، يلتزم العميل بسداد جميع المستحقات عن الفترة المنقضية حتى تاريخ الإنهاء الفعلي')

add_heading_styled('الملكية الفكرية', level=2)
add_bullet('جميع المحتوى والمواد الإبداعية المُنتجة خصيصاً للعميل تؤول ملكيتها بالكامل إلى العميل بعد سداد جميع المستحقات المالية')
add_bullet('تحتفظ بيراميديا بحقها في استخدام الأعمال المُنجزة ضمن معرض أعمالها (Portfolio) ما لم يُتفق على خلاف ذلك')

add_heading_styled('ملاحظات عامة', level=2)
add_bullet('الهوية البصرية: تم تسليمها مسبقاً بعقد منفصل')
add_bullet('الموقع الإلكتروني: مسؤولية العميل — تعمل بيراميديا على تحسين SEO فقط بعد تسليم الموقع')
add_bullet('منصة المحامين: مشروع منفصل بعقد مستقل')

# ==================== PAGE 10: SIGNATURE ====================
add_page_break()
add_heading_styled('التوقيع والاعتماد')
add_para('بتوقيع الطرفين أدناه، يُعتبر هذا العرض اتفاقاً ملزماً ونافذاً بالشروط والأحكام المذكورة أعلاه.', color=GRAY)

doc.add_paragraph()
add_heading_styled('الطرف الأول — بيراميديا', level=2)
for field in ['الاسم', 'المسمى الوظيفي', 'التوقيع', 'التاريخ']:
    add_para(f'{field}: _________________________________')

doc.add_paragraph()
add_heading_styled('الطرف الثاني — مركز إتمام للخدمات القضائية', level=2)
for field in ['الاسم', 'المسمى الوظيفي', 'رقم التواصل', 'التوقيع', 'التاريخ']:
    add_para(f'{field}: _________________________________')

doc.add_paragraph()
add_heading_styled('أشخاص اعتماد المحتوى (بحد أقصى شخصين)', level=2)
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
approval_headers = ['#', 'الاسم', 'رقم التواصل']
for i, h in enumerate(approval_headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'F97316')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(p)
        for r in p.runs:
            r.font.color.rgb = WHITE
            r.font.bold = True
for row_idx in [1, 2]:
    table.rows[row_idx].cells[0].text = str(row_idx)
    for p in table.rows[row_idx].cells[0].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_heading_styled('خيار الدفع المُختار', level=2)
for opt in ['☐  ربع سنوي (18,000 درهم/شهر)', '☐  نصف سنوي (16,000 درهم/شهر)', '☐  سنوي (15,000 درهم/شهر)']:
    add_para(opt, size=12)

doc.add_paragraph()
doc.add_paragraph()
add_para('سري وخاص — Pyramedia X © 2026', color=RGBColor(0xBB, 0xBB, 0xBB), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

# Save
out = '/home/node/openclaw/projects/etmam-marketing-sow/proposal.docx'
doc.save(out)
print(f'✅ DOCX saved: {out}')
